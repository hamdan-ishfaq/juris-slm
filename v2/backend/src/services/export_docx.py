"""Phase 9F — DOCX export for contract workspace."""
from __future__ import annotations

from io import BytesIO
from typing import Any


def build_docx_bytes(
    *,
    filename: str,
    content: str,
    clauses: list[dict[str, Any]] | None = None,
    annotations: list[dict[str, Any]] | None = None,
) -> bytes:
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    doc.add_heading(filename or "Contract", level=0)
    if clauses:
        doc.add_heading("Clause index", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Clause"
        table.rows[0].cells[1].text = "Title"
        for c in clauses:
            row = table.add_row().cells
            row[0].text = c.get("id", "")
            row[1].text = c.get("title", "")

    doc.add_heading("Body", level=1)
    for para in (content or "").split("\n"):
        p = doc.add_paragraph(para)
        p.style.font.size = Pt(11)

    if annotations:
        doc.add_heading("Annotations", level=1)
        for ann in annotations:
            doc.add_paragraph(f"{ann.get('clause_id', 'general')}: {ann.get('comment', '')}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
